"""Create a fragmented placeholder docx fixture for CI testing."""
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import os

doc = Document()

# Add a paragraph with fragmented {{ muvekkil_ad }} split across 4 <w:t> nodes
p = doc.add_paragraph()
r = p.add_run()

# The run's <w:r> element will contain multiple <w:t> children
# This simulates Word splitting a placeholder across XML nodes
rPr = r._r.get_or_add_rPr()

# Create 4 separate <w:t> elements inside the same <w:r>
parts = ['{{', ' muvekkil_', 'ad ', '}}']
for part in parts:
    t = etree.SubElement(r._r, qn('w:t'))
    t.text = part

# Also add {%p dosya_no%} as a normal (non-fragmented) placeholder
p2 = doc.add_paragraph()
p2.add_run('{%p dosya_no%}')

output_path = os.path.join(os.path.dirname(__file__), '..', 'tests', 'fixtures', 'test-template-fragmented.docx')
doc.save(output_path)
print(f'Created {output_path}')
