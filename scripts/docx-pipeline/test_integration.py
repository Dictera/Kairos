"""End-to-end integration test for docxtpl render + LibreOffice convert pipeline."""

from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

# Ensure sidecar modules are importable
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))


def _set_times_new_roman(run):
    """Apply Times New Roman font to a python-docx run (high Turkish char support)."""
    from docx.oxml.ns import qn

    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _build_canary_template(output_path: str) -> None:
    """Programmatically create a .docx fixture with Jinja2 tags and Times New Roman font."""
    from docx import Document

    doc = Document()

    paragraphs = [
        "Müvekkil: {{ muvekkil.ad | upper_tr }} {{ muvekkil.soyad | upper_tr }}",
        "Karşı taraf: {{ taraf.karsitaraf_ad | lower_tr }}",
        "Tutar: {{ dosya.talep_tutari | tr_currency }}",
        "Tarih: {{ dosya.kaza_tarihi | tarih }}",
        "{% if durusmalar %}Duruşma var{% else %}Duruşma yok{% endif %}",
        "{% for d in durusmalar %}{{ d.mahkeme_kurum }} — {{ d.tarih | tarih }}{% endfor %}",
        "çÇğĞıİöÖşŞüÜ İstanbul şirket müvekkil",
    ]

    for text in paragraphs:
        p = doc.add_paragraph()
        run = p.add_run(text)
        _set_times_new_roman(run)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


def _run_sidecar_command(envelope: dict) -> dict:
    """Pipe a JSON envelope to main.py and return the parsed JSON response."""
    main_py = os.path.join(os.path.dirname(__file__), "main.py")
    result = subprocess.run(
        [sys.executable, main_py],
        input=json.dumps(envelope),
        capture_output=True,
        text=True,
        timeout=60,
    )
    try:
        return json.loads(result.stdout)
    except json.JSONDecodeError as exc:
        raise AssertionError(
            f"Sidecar returned invalid JSON (rc={result.returncode}): {result.stdout!r}"
        ) from exc


def _get_libreoffice_path() -> str | None:
    """Detect LibreOffice soffice binary."""
    if env_path := os.environ.get("LIBREOFFICE_PATH"):
        return env_path
    for candidate in ["soffice", r"C:\Program Files\LibreOffice\program\soffice.exe"]:
        if shutil.which(candidate):
            return candidate
    return None


def test_integration() -> None:
    """Run full render → convert cycle and verify outputs."""
    fixture_dir = os.path.join(os.path.dirname(__file__), "fixtures")
    canary_path = os.path.join(fixture_dir, "canary-template.docx")
    temp_id = os.urandom(4).hex()
    rendered_path = os.path.join(tempfile.gettempdir(), f"canary-rendered-{temp_id}.docx")
    pdf_output_dir = tempfile.gettempdir()

    try:
        # 1. Build canary template
        _build_canary_template(canary_path)
        assert os.path.exists(canary_path), "Canary template was not created"

        # 2. Render
        context = {
            "muvekkil": {"ad": "ahmet", "soyad": "ışık"},
            "dosya": {"dosya_no": "2026/42", "talep_tutari": 150000, "kaza_tarihi": "2026-02-14"},
            "taraf": {"karsitaraf_ad": "MEHMET ÇELİK"},
            "durusmalar": [
                {"tarih": "2026-03-15", "mahkeme_kurum": "İstanbul Adliyesi"}
            ],
            "stk": {},
            "mahkeme": {},
            "sureler": [],
            "finans_kalemleri": [],
            "notlar": [],
        }

        render_envelope = {
            "command": "render",
            "params": {
                "template_path": canary_path,
                "output_path": rendered_path,
                "context": context,
            },
        }

        render_result = _run_sidecar_command(render_envelope)
        assert render_result.get("status") == "success", (
            f"Render failed: {render_result.get('message')}"
        )
        assert os.path.exists(rendered_path), "Rendered DOCX was not created"

        # 3. Convert (skip gracefully if LibreOffice is missing)
        libreoffice_path = _get_libreoffice_path()
        if not libreoffice_path:
            print(
                "WARNING: LibreOffice not found — skipping PDF conversion. "
                "Set LIBREOFFICE_PATH env var or install LibreOffice to run full pipeline."
            )
            return

        convert_envelope = {
            "command": "convert",
            "params": {
                "input_path": rendered_path,
                "output_dir": pdf_output_dir,
                "libreoffice_path": libreoffice_path,
                "timeout": 120,
            },
        }

        convert_result = _run_sidecar_command(convert_envelope)
        assert convert_result.get("status") == "success", (
            f"Convert failed: {convert_result.get('message')}"
        )

        pdf_path = convert_result["result"]["output_path"]
        assert os.path.exists(pdf_path), "PDF was not created"
        assert pdf_path.endswith(".pdf"), f"Output is not a PDF: {pdf_path}"
        assert os.path.getsize(pdf_path) > 0, "PDF is empty"

        # 4. Verify PDF text content (best-effort)
        _verify_pdf_text(pdf_path)

        print("Integration test passed!")

    finally:
        for p in (rendered_path,):
            if p and os.path.exists(p):
                os.unlink(p)


def _verify_pdf_text(pdf_path: str) -> None:
    """Extract text from PDF and assert expected content is present."""
    try:
        import PyPDF2
    except ImportError:
        try:
            import pdfplumber
        except ImportError:
            print(
                "WARNING: PyPDF2/pdfplumber not installed — skipping PDF text extraction. "
                "Install one to enable text assertions."
            )
            return

    text = ""
    try:
        if "PyPDF2" in sys.modules:
            with open(pdf_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
        else:
            import pdfplumber

            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
    except Exception as exc:
        print(f"WARNING: PDF text extraction failed: {exc}")
        return

    checks = [
        ("AHMET IŞIK", "upper_tr filter"),
        ("mehmet çelik", "lower_tr filter"),
        ("150.000,00", "tr_currency filter"),
        ("14.02.2026", "tarih filter"),
        ("Duruşma var", "conditional block"),
        ("İstanbul Adliyesi", "loop block"),
        ("çÇğĞıİöÖşŞüÜ", "Turkish characters"),
    ]

    missing = []
    for substring, desc in checks:
        if substring not in text:
            missing.append(f"  - {desc}: expected '{substring}'")

    if missing:
        raise AssertionError("PDF text verification failed:\n" + "\n".join(missing))


if __name__ == "__main__":
    test_integration()
